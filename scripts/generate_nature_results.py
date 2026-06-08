# generate_nature_results.py
import os
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def main():
    dest_path = "/mnt/c/Users/shiva/OneDrive/Desktop/PP_Nature_style_Results.docx"
    print(f"Generating Nature-style Results document at {dest_path}...")
    doc = Document()
    
    # Page setup
    teal_color = RGBColor(0, 90, 112)
    grey_color = RGBColor(128, 128, 128)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Document Title
    p_title = doc.add_paragraph()
    run = p_title.add_run("Genome-wide transcriptomic reorganization suggests resource allocation trade-offs and specialized chemical defense modulation in Glycyrrhiza uralensis under herbivore attack")
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = teal_color
    
    p_authors = doc.add_paragraph()
    run = p_authors.add_run("Shiva Raj Chippa\n")
    run.font.bold = True
    run.font.size = Pt(11)
    run = p_authors.add_run("Computational Biologist, IISER Tirupati Alumni\n")
    run.font.size = Pt(10)
    run.font.italic = True
    
    doc.add_paragraph()
    
    # 1. Abstract / Summary style paragraph
    p_lead = doc.add_paragraph()
    run = p_lead.add_run(
        "Plants respond to herbivory through complex networks of signal transduction and metabolic adjustments. "
        "Here, we profile the transcriptomic response of the medicinal legume Glycyrrhiza uralensis to chewing insect infestation "
        "across a 24-hour timecourse. Using strict homology filtering and statistical criteria, we show that G. uralensis "
        "reorganizes its transcriptome to suppress primary light-harvesting processes while activating flavonoid and "
        "phenylpropanoid biosynthesis. Our temporal analysis suggests a progressive signaling-to-metabolite response, "
        "identifying candidate novel species-specific defense genes and highlighting time-dependent jasmonate signaling dynamics."
    )
    run.font.size = Pt(11.5)
    run.font.bold = True
    
    # Section 1
    h1 = doc.add_heading("Global transcriptional dynamics and replicate validation", level=2)
    h1.runs[0].font.color.rgb = teal_color
    h1.runs[0].font.name = 'Times New Roman'
    
    doc.add_paragraph(
        "To understand G. uralensis transcriptional defense networks, we performed RNA-seq on leaves harvested at 0h, 6h, 12h, "
        "and 24h post-infestation. Principal Component Analysis (PCA) of VST-normalized counts revealed that biological replicates "
        "cluster tightly, indicating minimal technical variance. PC1 captures 56% of total variance and separates early stages (0-6h) "
        "from later responses (12-24h). Hierarchical distance heatmaps confirmed these groupings, showing that the transcriptome "
        "undergoes progressive, coordinated remodeling in response to insect chewing."
    )
    
    # Section 2
    h1 = doc.add_heading("Homology filtering isolates novel legume-specific candidates", level=2)
    h1.runs[0].font.color.rgb = teal_color
    h1.runs[0].font.name = 'Times New Roman'
    
    doc.add_paragraph(
        "To map GWH annotations to the model Arabidopsis thaliana proteome, we ran DIAMOND blastp. To prevent annotation over-claiming, "
        "homologs were filtered strictly: E-value <= 1e-5, identity >= 30%, and single best-hit selection based on bitscore. This logic "
        "resolved 25,121 high-confidence ortholog mappings (63.6% of GWH genes). Umapped DEGs represent candidate novel or lineage-specific "
        "genes: we identified 1,177 novel DEGs at 6h (13.64%), 1,015 at 12h (12.98%), and 1,309 at 24h (13.33%). These candidate novel genes "
        "represent a rich genomic resource for specialized defense features in Glycyrrhiza."
    )
    
    # Section 3
    h1 = doc.add_heading("Functional enrichment highlights growth-defense trade-offs", level=2)
    h1.runs[0].font.color.rgb = teal_color
    h1.runs[0].font.name = 'Times New Roman'
    
    doc.add_paragraph(
        "Over-representation analysis (ORA) using clusterProfiler (BH-adjusted padj < 0.05, qvalue < 0.05) separate upregulated "
        "and downregulated processes. Upregulated biological processes show robust induction of flavonoid biosynthesis (ath00941) "
        "and ribosome translation (ath03010) across all timepoints, driving chemical defense mechanisms. Downregulated pathways "
        "are enriched in photosynthesis thylakoid membrane proteins (ath00196) and light-harvesting antenna centers, suggesting "
        "suppression of photosynthesis to conserve energy resources (growth-defense trade-offs)."
    )
    
    # Section 4
    h1 = doc.add_heading("Time-dependent modulation of jasmonic and salicylic acid signaling", level=2)
    h1.runs[0].font.color.rgb = teal_color
    h1.runs[0].font.name = 'Times New Roman'
    
    doc.add_paragraph(
        "We observed time-dependent modulation of JA signaling components rather than simple pathway activation. Putative homologs "
        "of lipoxygenases (LOX4) and allene oxide cyclase (AOC1) were differentially regulated. Putative master transcription factor MYC2 "
        "showed a delayed induction, remaining non-significant early (6h, padj = 0.18) but showing strong upregulation later (12h, 24h). "
        "Conversely, repressor JAZ1 homologs exhibited dynamic regulatory behavior (downregulated at 6h, upregulated at 12h, and downregulated at 24h). "
        "Salicylic acid (SA) homologs (SARD1, TGA3) showed limited evidence of SA-associated transcriptional modulation under chewing stress, "
        "while PR1 was unmapped and NPR1 was filtered due to low counts."
    )
    
    doc.add_paragraph(
        "Upregulation of genes associated with secondary metabolism (PAL1, C4H, 4CL1, CHS, CHI) suggests potential accumulation "
        "of protective flavonoids (liquiritin/liquiritigenin precursors), although metabolite concentrations were not directly measured."
    )
    
    # Add Table 2: Representative Genes
    table2_csv_path = "/home/shiva/bioinformatics_plant_projects/licorice_insect_rnaseq/results/table2_representative_defense_genes.csv"
    if os.path.exists(table2_csv_path):
        table2_rows = []
        with open(table2_csv_path, 'r') as f:
            reader = csv.reader(f)
            header = next(reader)
            for row in reader:
                if row:
                    table2_rows.append(row)
        
        doc.add_paragraph("\nTable 1: Expression fold changes (log2FC) and BH-adjusted p-values (padj) for key representative plant defense genes.")
        
        table2 = doc.add_table(rows=len(table2_rows) + 1, cols=6)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        table2.style = 'Table Grid'
        
        col_widths2 = [Inches(0.9), Inches(1.2), Inches(1.6), Inches(1.1), Inches(1.1), Inches(1.1)]
        for row in table2.rows:
            for idx, width in enumerate(col_widths2):
                row.cells[idx].width = width
        
        t2_headers = ["Symbol", "G. uralensis Locus", "Functional Group", "6H log2FC\n(padj)", "12H log2FC\n(padj)", "24H log2FC\n(padj)"]
        t2_hdr_cells = table2.rows[0].cells
        for i, h_text in enumerate(t2_headers):
            t2_hdr_cells[i].text = h_text
            t2_hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            t2_hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = teal_color
            t2_hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9.5)
            
        def format_padj(val_str):
            if val_str == "N/A" or not val_str:
                return "N/A*"
            try:
                val = float(val_str)
                if val == 1.0 or val == 0.0:
                    return f"{val:.1f}"
                return f"{val:.2e}"
            except ValueError:
                return "N/A*"

        def format_lfc(val_str):
            if val_str == "N/A" or not val_str:
                return "N/A"
            try:
                val = float(val_str)
                return f"{val:+.2f}"
            except ValueError:
                return "N/A"

        for r_idx, row_data in enumerate(table2_rows):
            at_id, gwh_id, symbol, group, desc, lfc_6h, padj_6h, lfc_12h, padj_12h, lfc_24h, padj_24h = row_data
            
            row_cells = table2.rows[r_idx + 1].cells
            row_cells[0].text = symbol
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[1].text = gwh_id
            row_cells[2].text = group
            
            row_cells[3].text = f"{format_lfc(lfc_6h)}\n({format_padj(padj_6h)})" if lfc_6h != "N/A" else "N/A"
            row_cells[4].text = f"{format_lfc(lfc_12h)}\n({format_padj(padj_12h)})" if lfc_12h != "N/A" else "N/A"
            row_cells[5].text = f"{format_lfc(lfc_24h)}\n({format_padj(padj_24h)})" if lfc_24h != "N/A" else "N/A"
            
            for cell in row_cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(9.0)
                        
        p_note = doc.add_paragraph()
        run = p_note.add_run("*N/A represents genes filtered out by DESeq2 independent filtering or genes with missing values due to mapping annotation limitations.")
        run.font.size = Pt(8.5)
        run.font.italic = True
        
    doc.add_page_break()
    
    # Section 5
    h1 = doc.add_heading("Progressive transcriptomic response model", level=2)
    h1.runs[0].font.color.rgb = teal_color
    h1.runs[0].font.name = 'Times New Roman'
    
    doc.add_paragraph(
        "In summary, our temporal transcriptomics suggest a progressive response cascade. Early insect chewing is associated with cell-wall wounding, "
        "triggering Ca2+ and MAPK signaling cascades at 6h. This signaling coincides with a time-dependent phytohormone modulation (JA pathway) at 12h, "
        "allowing transcription factor activation (such as MYC2 and WRKYs). By 24h, this transcription program contributes to downstream defenses, "
        "including upregulated flavonoid synthesis and redox homeostasis enzymes (SOD, GSTs), while growth-associated photosynthesis genes are repressed."
    )
    
    doc.save(dest_path)
    print("Nature-style Results document generated successfully!")

if __name__ == "__main__":
    main()
